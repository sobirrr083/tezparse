#!/usr/bin/env python3
"""
Quiz Format Converter  (+++/===/# → A/B/C/D + ANSWER: X)
----------------------------------------------------------
Kirish format belgilari:
  +++  (3+ ta '+') — yangi savol boshlanishi; savol matni shu qatorda yoki
                      keyingi qatorda bo'lishi mumkin.
  ===  (3+ ta '=') — variant ajratgichi / chegara qatori.
  #    — variant boshida kelsa, shu variant TO'G'RI javob.

Qo'llab-quvvatlanadigan fayl turlari:
  .txt  .docx  .doc  .pdf  .xlsx  .xls

Ishlatish:
  python quiz_converter_new.py input.txt
  python quiz_converter_new.py input.docx output.txt
  python quiz_converter_new.py input.pdf --non-interactive
"""

import re
import sys
import os
from pathlib import Path

LETTERS = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"

# ─────────────────────────────────────────────
# 1. FAYL O'QISH  (asl converterdan olingan)
# ─────────────────────────────────────────────

def extract_text(filepath: str) -> str:
    ext = Path(filepath).suffix.lower()

    if ext == ".txt":
        for enc in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
            try:
                with open(filepath, encoding=enc) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        with open(filepath, encoding="utf-8", errors="replace") as f:
            return f.read()

    elif ext == ".docx":
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx o'rnatilmagan: pip install python-docx")
        doc = Document(filepath)
        parts = []
        for para in doc.paragraphs:
            parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)

    elif ext == ".doc":
        import subprocess
        try:
            result = subprocess.run(
                ["antiword", filepath], capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0:
                return result.stdout
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
        try:
            import textract
            return textract.process(filepath).decode("utf-8", errors="replace")
        except ImportError:
            raise ImportError(
                ".doc uchun antiword yoki textract kerak.\n"
                "  sudo apt install antiword  yoki  pip install textract"
            )

    elif ext == ".pdf":
        try:
            import pdfplumber
            pages = []
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        pages.append(t)
            return "\n".join(pages)
        except ImportError:
            pass
        try:
            import PyPDF2
            pages = []
            with open(filepath, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    pages.append(page.extract_text() or "")
            return "\n".join(pages)
        except ImportError:
            raise ImportError(
                "PDF uchun pdfplumber yoki PyPDF2 kerak.\n"
                "  pip install pdfplumber"
            )

    elif ext == ".xlsx":
        try:
            import openpyxl
        except ImportError:
            raise ImportError("openpyxl o'rnatilmagan: pip install openpyxl")
        wb = openpyxl.load_workbook(filepath, data_only=True)
        lines = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                row_text = " ".join(
                    str(c).strip() for c in row if c is not None and str(c).strip()
                )
                if row_text:
                    lines.append(row_text)
        return "\n".join(lines)

    elif ext == ".xls":
        try:
            import xlrd
        except ImportError:
            raise ImportError("xlrd o'rnatilmagan: pip install xlrd")
        wb = xlrd.open_workbook(filepath)
        lines = []
        for sheet in wb.sheets():
            for i in range(sheet.nrows):
                row_text = " ".join(
                    str(v).strip() for v in sheet.row_values(i) if str(v).strip()
                )
                if row_text:
                    lines.append(row_text)
        return "\n".join(lines)

    else:
        raise ValueError(f"Qo'llab-quvvatlanmagan format: {ext}")


# ─────────────────────────────────────────────
# 2. YANGI FORMAT PARSE QILISH
# ─────────────────────────────────────────────

# +++ (3 yoki undan ko'p '+') — savol separator
PAT_Q_SEP  = re.compile(r"^\+{3,}")
# === (3 yoki undan ko'p '=') — variant separator
PAT_V_SEP  = re.compile(r"^={3,}")
# Variant boshida '#' belgisi — to'g'ri javob
PAT_CORRECT_MARK = re.compile(r"^#\s*")


def _clean(text: str) -> str:
    """Matndan ortiqcha bo'sh joylarni tozalash."""
    return re.sub(r"\s+", " ", text).strip()


def parse_blocks(text: str) -> list:
    """
    Matnni savol bloklarga ajratadi.

    Mantiq:
      +++... qatori → yangi savol boshlanadi.
              Agar +++ bilan birga matn bo'lsa, u savol matni.
              Agar +++ yolg'iz bo'lsa, keyingi qator savol matni.
      ===... qatori → variant ajratgichi (chegaralovchi), o'zi hech narsa emas.
              Undan keyingi matn — yangi variant (variant bo'shliq matnlari
              keyingi === yoki +++ kelguncha to'planadi).
      # boshidagi variant — to'g'ri javob.
    """
    lines = text.splitlines()
    blocks = []

    # Holat mashina
    STATE_BEFORE_Q   = "before_q"    # savol hali boshlanmagan
    STATE_IN_Q       = "in_q"        # savol matni to'planmoqda
    STATE_IN_VARIANT = "in_variant"  # variant matni to'planmoqda

    state = STATE_BEFORE_Q
    current = None          # joriy blok dict
    variant_buf = []        # joriy variant qatorlari
    variant_correct = False # joriy variant to'g'rimi

    def _flush_variant():
        """variant_buf ni current['variants'] ga qo'shish."""
        nonlocal variant_buf, variant_correct
        vtext = _clean(" ".join(variant_buf))
        if vtext:
            current["variants"].append({
                "text": vtext,
                "correct": variant_correct,
            })
        variant_buf = []
        variant_correct = False

    def _flush_block():
        """current blokni blocks ga qo'shish."""
        if current is not None:
            if state == STATE_IN_VARIANT:
                _flush_variant()
            blocks.append(current)

    for raw in lines:
        line = raw.strip()

        # +++ separator
        if PAT_Q_SEP.match(line):
            # Oldingi blokni yopish
            if current is not None:
                if state == STATE_IN_VARIANT:
                    _flush_variant()
                blocks.append(current)

            # Savol matni shu qatorda bo'lishi mumkin (masalan: "+++Savol matni")
            inline_q = PAT_Q_SEP.sub("", line).strip()
            current = {"question": inline_q, "variants": []}
            variant_buf = []
            variant_correct = False
            state = STATE_IN_Q if not inline_q else STATE_IN_Q
            continue

        # === separator
        if PAT_V_SEP.match(line):
            if current is None:
                # Hali savol yo'q, o'tkazib yuborish
                continue
            if state == STATE_IN_VARIANT:
                _flush_variant()
            # Keyingi qator yangi variant bo'ladi
            state = STATE_IN_VARIANT
            continue

        # Oddiy matn qatori
        if current is None:
            # Savol boshlanmagan — o'tkazib yuborish
            continue

        if state == STATE_IN_Q:
            # Savol matni to'planmoqda
            if line:
                if current["question"]:
                    current["question"] += " " + line
                else:
                    current["question"] = line
            continue

        if state == STATE_IN_VARIANT:
            if not line:
                continue
            # Birinchi qator bo'lsa — to'g'ri/noto'g'ri aniqlash
            if not variant_buf:
                if PAT_CORRECT_MARK.match(line):
                    variant_correct = True
                    line = PAT_CORRECT_MARK.sub("", line).strip()
            variant_buf.append(line)
            continue

    # Oxirgi blokni yopish
    if current is not None:
        if state == STATE_IN_VARIANT:
            _flush_variant()
        blocks.append(current)

    return blocks


# ─────────────────────────────────────────────
# 3. VALIDATSIYA
# ─────────────────────────────────────────────

def validate_block(block: dict) -> list:
    issues = []
    correct = [v for v in block["variants"] if v["correct"]]
    wrong   = [v for v in block["variants"] if not v["correct"]]

    if not block["question"].strip():
        issues.append("empty_question")
    if len(block["variants"]) < 2:
        issues.append("too_few_variants")
    if len(correct) == 0:
        issues.append("no_correct")
    if len(correct) > 1:
        issues.append("multiple_correct")
    if len(wrong) == 0:
        issues.append("no_wrong")
    return issues


# ─────────────────────────────────────────────
# 4. FORMATLASH
# ─────────────────────────────────────────────

def format_block(block: dict, number: int) -> str:
    lines = []
    lines.append(f"{number}. {block['question'].strip()}")
    lines.append("")

    correct_letter = None
    for i, variant in enumerate(block["variants"]):
        letter = LETTERS[i] if i < len(LETTERS) else f"?{i}"
        lines.append(f"    {letter}. {variant['text']}")
        if variant["correct"]:
            correct_letter = letter

    lines.append("")
    lines.append(f"    ANSWER: {correct_letter or '?'}")
    lines.append("")
    lines.append("")
    return "\n".join(lines)


# ─────────────────────────────────────────────
# 5. INTERAKTIV TUZATISH
# ─────────────────────────────────────────────

def fix_block_interactive(block: dict, idx: int) -> bool:
    issues = validate_block(block)
    print(f"\n{'─'*60}")
    print(f"  [Blok {idx+1}] Muammo: {', '.join(issues)}")
    print(f"  Savol: {block['question'][:80]}")
    print(f"  Variantlar:")
    for j, v in enumerate(block["variants"]):
        marker = "#" if v["correct"] else " "
        print(f"    {j+1}. [{marker}] {v['text']}")
    print(f"{'─'*60}")

    if "no_correct" in issues or "multiple_correct" in issues:
        print("  Qaysi variant to'g'ri? Raqamini kiriting, yoki 's' — o'tkazib yuborish:")
        for j, v in enumerate(block["variants"]):
            print(f"    {j+1}. {v['text']}")
        choice = input("  > ").strip()
        if choice.lower() == "s":
            return False
        try:
            chosen = int(choice) - 1
            if 0 <= chosen < len(block["variants"]):
                for j, v in enumerate(block["variants"]):
                    v["correct"] = (j == chosen)
            else:
                print("  Noto'g'ri raqam. O'tkazib yuborildi.")
                return False
        except ValueError:
            print("  Noto'g'ri kiritma. O'tkazib yuborildi.")
            return False

    if "too_few_variants" in issues or "empty_question" in issues:
        print("  Blok to'liq emas, o'tkazib yuborildi.")
        return False

    return True


# ─────────────────────────────────────────────
# 6. ASOSIY FUNKSIYA
# ─────────────────────────────────────────────

def convert(input_path: str, output_path: str = None, interactive: bool = True) -> tuple:
    print(f"\n📂 O'qilmoqda: {input_path}")
    text = extract_text(input_path)

    blocks = parse_blocks(text)
    print(f"🔍 Topildi: {len(blocks)} ta savol bloki")

    output_parts = []
    errors = []
    q_num = 1

    for i, block in enumerate(blocks):
        issues = validate_block(block)

        if issues:
            if interactive:
                ok = fix_block_interactive(block, i)
                if not ok:
                    errors.append(
                        f"Blok {i+1} o'tkazildi | {issues} | {block['question'][:50]}"
                    )
                    continue
                remaining = validate_block(block)
                if remaining:
                    errors.append(f"Blok {i+1} hali ham xatolikli: {remaining}")
                    continue
            else:
                errors.append(
                    f"Blok {i+1} o'tkazildi | {issues} | {block['question'][:50]}"
                )
                continue

        output_parts.append(format_block(block, q_num))
        q_num += 1

    result_text = "\n".join(output_parts)

    if not output_path:
        stem   = Path(input_path).stem
        parent = Path(input_path).parent
        output_path = str(parent / f"{stem}_converted.txt")

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(result_text)

    converted = q_num - 1
    skipped   = len(errors)

    print(f"\n✅ Konvertatsiya yakunlandi!")
    print(f"   Saqlandi  : {output_path}")
    print(f"   Savollar  : {converted} ta")
    print(f"   O'tkazildi: {skipped} ta")

    if errors:
        print(f"\n⚠️  Ogohlantirmalar ({len(errors)}):")
        for e in errors:
            print(f"   • {e}")

    return output_path, skipped, errors


# ─────────────────────────────────────────────
# 7. CLI
# ─────────────────────────────────────────────

def main():
    args = sys.argv[1:]

    if not args or args[0] in ("-h", "--help"):
        print(__doc__)
        sys.exit(0)

    non_interactive = "--non-interactive" in args
    args = [a for a in args if not a.startswith("--")]

    input_file  = args[0]
    output_file = args[1] if len(args) > 1 else None

    if not os.path.isfile(input_file):
        print(f"❌ Fayl topilmadi: {input_file}")
        sys.exit(1)

    try:
        convert(input_file, output_file, interactive=not non_interactive)
    except ImportError as e:
        print(f"\n❌ Kutubxona xatosi: {e}")
        sys.exit(2)
    except Exception as e:
        print(f"\n❌ Xato: {e}")
        raise


if __name__ == "__main__":
    main()
      
