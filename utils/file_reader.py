"""
utils/file_reader.py
---------------------
Fayl turlaridan matn chiqarish.
"""

from pathlib import Path

SUPPORTED_EXTENSIONS = {".txt", ".docx", ".doc", ".pdf", ".xlsx", ".xls"}


def extract_text(filepath: str) -> str:
    ext = Path(filepath).suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Qollab-quvvatlanmagan format: '{ext}'\n"
            f"Mumkin bolgan formatlar: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    if ext == ".txt":
        return _read_txt(filepath)
    elif ext == ".docx":
        return _read_docx(filepath)
    elif ext == ".doc":
        return _read_doc(filepath)
    elif ext == ".pdf":
        return _read_pdf(filepath)
    elif ext == ".xlsx":
        return _read_xlsx(filepath)
    elif ext == ".xls":
        return _read_xls(filepath)


def _read_txt(filepath):
    for enc in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            with open(filepath, encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    with open(filepath, encoding="utf-8", errors="replace") as f:
        return f.read()


def _read_docx(filepath):
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx kerak: pip install python-docx")
    doc = Document(filepath)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _read_doc(filepath):
    import subprocess
    try:
        r = subprocess.run(
            ["antiword", filepath], capture_output=True, text=True, timeout=30
        )
        if r.returncode == 0:
            return r.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    try:
        import textract
        return textract.process(filepath).decode("utf-8", errors="replace")
    except ImportError:
        raise ImportError(".doc uchun antiword yoki textract kerak.")


def _read_pdf(filepath):
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    pages.append(t)
        return "\n".join(pages)
    except ImportError:
        pass
    try:
        import PyPDF2
        pages = []
        with open(filepath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                pages.append(page.extract_text() or "")
        return "\n".join(pages)
    except ImportError:
        raise ImportError("PDF uchun: pip install pdfplumber")


def _read_xlsx(filepath):
    try:
        import openpyxl
    except ImportError:
        raise ImportError("openpyxl kerak: pip install openpyxl")
    wb = openpyxl.load_workbook(filepath, data_only=True)
    lines = []
    for sheet in wb.worksheets:
        for row in sheet.iter_rows(values_only=True):
            row_text = " ".join(
                str(c).strip() for c in row if c is not None and str(c).strip()
            )
            if row_text:
                lines.append(row_text)
    return "\n".join(lines)


def _read_xls(filepath):
    try:
        import xlrd
    except ImportError:
        raise ImportError("xlrd kerak: pip install xlrd")
    wb = xlrd.open_workbook(filepath)
    lines = []
    for sheet in wb.sheets():
        for i in range(sheet.nrows):
            row_text = " ".join(
                str(v).strip() for v in sheet.row_values(i) if str(v).strip()
            )
            if row_text:
                lines.append(row_text)
    return "\n".join(lines)
